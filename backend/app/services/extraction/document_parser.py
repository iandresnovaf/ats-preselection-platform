"""Parser de documentos - Extrae texto de PDF/DOCX y detecta tipo de documento."""
import io
import logging
import os
import re
import tempfile
import hashlib
import time
from pathlib import Path
from typing import Optional, Dict, Any, List
from datetime import datetime

import pdfplumber
from docx import Document as DocxDocument

from app.core.config import settings
from app.services.extraction.models import (
    DocumentType, ProcessingStatus, ParseResult
)

logger = logging.getLogger(__name__)


class DocumentParserError(Exception):
    """Error en parsing de documentos."""
    pass


class DocumentNotFoundError(DocumentParserError):
    """Documento no encontrado."""
    pass


class UnsupportedDocumentError(DocumentParserError):
    """Tipo de documento no soportado."""
    pass


class DocumentParser:
    """Parser de documentos para extracción de texto y detección de tipo."""
    
    # Extensiones soportadas
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.rtf'}
    
    # Patrones para detección de tipo de documento
    ASSESSMENT_PATTERNS = [
        r'\b(factor\s+oscuro|dark\s+factor|dfi)\b',
        r'\b(inteligencia\s+ejecutiva|executive\s+intelligence)\b',
        r'\b(psicométrica|psychometric|assessment|evaluación)\b',
        r'\b(egocentrismo|egocentrism|narcisismo|narcissism)\b',
        r'\b(volatilidad|volatility|sicopatía|psychopathy)\b',
        r'\b(manuabilidad|manipulativeness|machiavellian)\b',
        r'\b(sinceridad|sincerity|consistencia|consistency)\b',
        r'\bpuntaje\s*total|total\s*score\b',
        r'\bdimensión|dimension\s+score\b',
        r'\bpercentil|percentile\b',
        r'\bperfil\s+gráfico|graphic\s+profile\b',
    ]
    
    INTERVIEW_PATTERNS = [
        r'\b(entrevista|interview|conversación|conversation)\b',
        r'\b(entrevistador|interviewer|entrevistado|interviewee)\b',
        r'\b(pregunta|question|respuesta|answer)\b',
        r'\b(evaluación\s+de\s+competencias|competency\s+assessment)\b',
        r'\b(soft\s+skills|competencias\s+blandas)\b',
        r'\b(recomendación|recommendation|conclusión|conclusion)\b',
        r'\b(situación|situational|behavioral|conductual)\b',
    ]
    
    CV_PATTERNS = [
        r'\b(curriculum|cv|resume|hoja\s+de\s+vida)\b',
        r'\b(experiencia\s+laboral|work\s+experience|employment)\b',
        r'\b(educación|education|formación|training)\b',
        r'\b(habilidades|skills|competencias|competencies)\b',
        r'\b(objetivo|objective|perfil|profile|summary)\b',
        r'\b(referencias|references)\b',
        r'\b(proyectos|projects|publicaciones|publications)\b',
    ]
    
    def __init__(self, storage_path: Optional[str] = None):
        """Inicializa el parser.
        
        Args:
            storage_path: Ruta base para almacenamiento de documentos
        """
        self.storage_path = storage_path or settings.UPLOAD_DIR
        self._ocr_engine = None
    
    async def parse_document(self, document_id: str, file_path: str, 
                            mime_type: Optional[str] = None) -> ParseResult:
        """Parsea un documento y extrae su texto.
        
        Args:
            document_id: ID del documento
            file_path: Ruta al archivo
            mime_type: MIME type del documento (opcional)
            
        Returns:
            ParseResult con el texto extraído y metadatos
        """
        start_time = time.time()
        
        try:
            # Verificar que el archivo existe
            if not os.path.exists(file_path):
                raise DocumentNotFoundError(f"Archivo no encontrado: {file_path}")
            
            # Detectar MIME type si no se proporciona
            if not mime_type:
                mime_type = self._detect_mime_type(file_path)
            
            # Extraer texto según el tipo
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.pdf':
                text, metadata = await self._extract_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                text, metadata = await self._extract_docx(file_path)
            elif file_ext == '.txt':
                text, metadata = await self._extract_txt(file_path)
            else:
                raise UnsupportedDocumentError(f"Tipo de archivo no soportado: {file_ext}")
            
            # Calcular hash del contenido
            content_hash = hashlib.sha256(text.encode('utf-8')).hexdigest()
            metadata['content_hash'] = content_hash
            
            # Detectar tipo de documento
            doc_type = self._detect_document_type(text)
            metadata['detected_type'] = doc_type.value
            
            processing_time = int((time.time() - start_time) * 1000)
            
            logger.info(
                f"Documento {document_id} parseado: "
                f"tipo={doc_type.value}, "
                f"caracteres={len(text)}, "
                f"tiempo={processing_time}ms"
            )
            
            return ParseResult(
                document_id=document_id,
                status=ProcessingStatus.PARSING,
                text=text,
                document_type=doc_type,
                metadata=metadata,
                processing_time_ms=processing_time
            )
            
        except Exception as e:
            logger.error(f"Error parseando documento {document_id}: {e}")
            return ParseResult(
                document_id=document_id,
                status=ProcessingStatus.ERROR,
                text="",
                document_type=DocumentType.OTHER,
                error_message=str(e),
                processing_time_ms=int((time.time() - start_time) * 1000)
            )
    
    def _detect_mime_type(self, file_path: str) -> str:
        """Detecta el MIME type de un archivo.
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            MIME type detectado
        """
        import mimetypes
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type or 'application/octet-stream'
    
    async def _extract_pdf(self, file_path: str) -> tuple:
        """Extrae texto de un PDF.
        
        Args:
            file_path: Ruta al PDF
            
        Returns:
            Tuple (texto_extraído, metadatos)
        """
        text_parts = []
        metadata = {
            'pages': 0,
            'has_images': False,
            'extraction_method': 'pdfplumber'
        }
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for i, page in enumerate(pdf.pages):
                    # Extraer texto
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    
                    # Detectar si tiene imágenes
                    if page.images:
                        metadata['has_images'] = True
                    
                    # Extraer metadatos de la primera página
                    if i == 0:
                        meta = pdf.metadata
                        if meta:
                            metadata.update({
                                'author': meta.get('Author'),
                                'creator': meta.get('Creator'),
                                'producer': meta.get('Producer'),
                                'creation_date': meta.get('CreationDate'),
                            })
                
                # Si hay muy poco texto, podría ser un PDF escaneado
                full_text = '\n\n'.join(text_parts)
                if len(full_text.strip()) < 100:
                    metadata['likely_scanned'] = True
                    logger.warning(f"PDF posiblemente escaneado detectado: {file_path}")
                
        except Exception as e:
            logger.error(f"Error extrayendo PDF: {e}")
            raise
        
        return '\n\n'.join(text_parts), metadata
    
    async def _extract_docx(self, file_path: str) -> tuple:
        """Extrae texto de un documento Word.
        
        Args:
            file_path: Ruta al DOCX
            
        Returns:
            Tuple (texto_extraído, metadatos)
        """
        text_parts = []
        metadata = {
            'extraction_method': 'python-docx',
            'has_tables': False,
            'paragraphs': 0
        }
        
        try:
            doc = DocxDocument(file_path)
            
            # Extraer propiedades del documento
            if doc.core_properties:
                metadata.update({
                    'author': doc.core_properties.author,
                    'title': doc.core_properties.title,
                    'created': doc.core_properties.created,
                    'modified': doc.core_properties.modified,
                })
            
            # Extraer texto de párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            metadata['paragraphs'] = len(doc.paragraphs)
            
            # Extraer texto de tablas (importante para assessments)
            if doc.tables:
                metadata['has_tables'] = True
                metadata['table_count'] = len(doc.tables)
                
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_text.append(' | '.join(row_text))
                    
                    if table_text:
                        text_parts.append('\n'.join(table_text))
        
        except Exception as e:
            logger.error(f"Error extrayendo DOCX: {e}")
            raise
        
        return '\n\n'.join(text_parts), metadata
    
    async def _extract_txt(self, file_path: str) -> tuple:
        """Extrae texto de un archivo de texto plano.
        
        Args:
            file_path: Ruta al TXT
            
        Returns:
            Tuple (texto_extraído, metadatos)
        """
        metadata = {
            'extraction_method': 'text_reader'
        }
        
        try:
            # Detectar encoding
            import chardet
            
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                encoding = detected['encoding'] or 'utf-8'
                metadata['encoding'] = encoding
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                text = f.read()
        
        except Exception as e:
            logger.error(f"Error extrayendo TXT: {e}")
            raise
        
        return text, metadata
    
    def _detect_document_type(self, text: str) -> DocumentType:
        """Detecta el tipo de documento basado en el contenido.
        
        Args:
            text: Texto extraído del documento
            
        Returns:
            DocumentType detectado
        """
        if not text:
            return DocumentType.OTHER
        
        text_lower = text.lower()
        scores = {
            DocumentType.ASSESSMENT: 0,
            DocumentType.INTERVIEW: 0,
            DocumentType.CV: 0,
        }
        
        # Contar coincidencias para assessments
        for pattern in self.ASSESSMENT_PATTERNS:
            matches = len(re.findall(pattern, text_lower, re.IGNORECASE))
            scores[DocumentType.ASSESSMENT] += matches
        
        # Contar coincidencias para entrevistas
        for pattern in self.INTERVIEW_PATTERNS:
            matches = len(re.findall(pattern, text_lower, re.IGNORECASE))
            scores[DocumentType.INTERVIEW] += matches
        
        # Contar coincidencias para CVs
        for pattern in self.CV_PATTERNS:
            matches = len(re.findall(pattern, text_lower, re.IGNORECASE))
            scores[DocumentType.CV] += matches
        
        # Determinar el tipo con mayor score
        max_type = max(scores, key=scores.get)
        max_score = scores[max_type]
        
        # Si no hay suficientes coincidencias, considerar como OTHER
        if max_score < 2:
            return DocumentType.OTHER
        
        # Si hay empate o scores muy cercanos, usar heurísticas adicionales
        if scores[DocumentType.ASSESSMENT] >= 3:
            return DocumentType.ASSESSMENT
        elif scores[DocumentType.INTERVIEW] >= 3:
            return DocumentType.INTERVIEW
        elif scores[DocumentType.CV] >= 2:
            return DocumentType.CV
        
        return DocumentType.OTHER
    
    async def parse_from_bytes(self, document_id: str, file_bytes: bytes, 
                              filename: str) -> ParseResult:
        """Parsea un documento desde bytes.
        
        Args:
            document_id: ID del documento
            file_bytes: Contenido del archivo
            filename: Nombre del archivo
            
        Returns:
            ParseResult con el texto extraído
        """
        # Guardar en archivo temporal
        suffix = Path(filename).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            result = await self.parse_document(document_id, tmp_path)
            return result
        finally:
            os.unlink(tmp_path)
    
    def calculate_file_hash(self, file_path: str) -> str:
        """Calcula el hash SHA-256 de un archivo.
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            Hash SHA-256
        """
        sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                sha256.update(chunk)
        return sha256.hexdigest()
    
    async def dedupe_by_hash(self, sha256_hash: str, db_session) -> Optional[str]:
        """Busca si ya existe un documento con el mismo hash.
        
        Args:
            sha256_hash: Hash a buscar
            db_session: Sesión de base de datos
            
        Returns:
            ID del documento existente o None
        """
        from sqlalchemy import select
        from app.models.rhtools import Document
        
        result = await db_session.execute(
            select(Document).where(Document.checksum == sha256_hash)
        )
        existing = result.scalar_one_or_none()
        
        if existing:
            return str(existing.id)
        
        return None
