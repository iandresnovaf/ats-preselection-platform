"""
Job Profile Extractor Service
Extrae información estructurada de perfiles de cargo desde documentos PDF/Word.
"""
import re
from typing import Dict, List, Optional, Any
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class JobProfileExtractor:
    """Extrae información estructurada de perfiles de cargo en PDF/Word."""
    
    SECTIONS = {
        "objetivo": ["objetivo del rol", "propósito", "misión", "objetivo", "purpose"],
        "estructura": ["estructura del cargo", "organización", "jerarquía", "structure", "organigrama"],
        "responsabilidades": ["responsabilidades", "funciones", "tareas", "responsibilities", "functions"],
        "perfil": ["perfil requerido", "requisitos", "perfil del candidato", "profile", "requirements"],
        "formacion": ["formación", "educación", "estudios", "education", "academic"],
        "experiencia": ["experiencia", "tiempo de experiencia", "experience"],
        "conocimientos": ["conocimientos técnicos", "skills técnicos", "competencias técnicas", "technical skills", "conocimientos"],
        "competencias": ["competencias clave", "habilidades", "soft skills", "competencies", "skills"],
        "disc": ["perfil disc", "comportamiento", "estilo", "disc", "behavioral profile"],
        "herramientas": ["herramientas", "tools", "software", "aplicaciones", "programs"],
        "condiciones": ["condiciones", "salario", "beneficios", "ubicación", "conditions", "compensation"]
    }
    
    # Patrones para extraer datos estructurados de la sección "Estructura"
    HIERARCHY_PATTERNS = {
        "reports_to": [
            r"reporta[\s\w]*:\s*([^\n]+)",
            r"reporta\s+a:\s*([^\n]+)",
            r"reports?\s+to:\s*([^\n]+)",
        ],
        "direct_reports": [
            r"personas?\s+a\s+cargo:\s*([^\n]+)",
            r"equipo\s+a\s+cargo:\s*([^\n]+)",
            r"direct\s+reports?:\s*([^\n]+)",
            r"reports?\s+directly:\s*([^\n]+)",
        ],
        "level": [
            r"nivel:\s*([^\n]+)",
            r"level:\s*([^\n]+)",
            r"seniority:\s*([^\n]+)",
        ],
        "work_mode": [
            r"modalidad:\s*([^\n]+)",
            r"modo\s+de\s+trabajo:\s*([^\n]+)",
            r"work\s+mode:\s*([^\n]+)",
            r"modalidad\s+laboral:\s*([^\n]+)",
        ],
        "location": [
            r"ubicaci[oó]n:\s*([^\n]+)",
            r"lugar:\s*([^\n]+)",
            r"location:\s*([^\n]+)",
        ]
    }
    
    # Patrones para requisitos
    REQUIREMENTS_PATTERNS = {
        "education": [
            r"formaci[oó]n:\s*([^\n]+(?:\n(?![A-Z][a-z]+:)[^\n]+)*)",
            r"educaci[oó]n:\s*([^\n]+(?:\n(?![A-Z][a-z]+:)[^\n]+)*)",
            r"estudios:\s*([^\n]+(?:\n(?![A-Z][a-z]+:)[^\n]+)*)",
        ],
        "experience_years": [
            r"tiempo\s+de\s+experiencia:\s*([^\n]+)",
            r"experiencia:\s*m[ií]nimo\s*([^\n]+)",
            r"años\s+de\s+experiencia:\s*([^\n]+)",
        ],
    }
    
    def __init__(self):
        self.errors = []
        self.warnings = []
    
    def extract_from_document(self, file_path: str) -> Dict[str, Any]:
        """
        Extrae información de perfil de cargo desde un archivo PDF, Word o TXT.
        
        Args:
            file_path: Ruta al archivo (PDF, DOCX, DOC, TXT)
            
        Returns:
            Dict con los datos estructurados del perfil
        """
        self.errors = []
        self.warnings = []
        
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        # Extraer texto según el formato
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            text = self._extract_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            text = self._extract_from_docx(file_path)
        elif extension in ['.txt', '.md']:
            text = self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Formato no soportado: {extension}")
        
        if not text or not text.strip():
            raise ValueError("No se pudo extraer texto del documento. Puede ser un documento escaneado.")
        
        # Parsear el contenido
        return self.parse_job_profile(text)
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extraer texto de un archivo PDF."""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extrayendo PDF: {e}")
            self.errors.append(f"Error al procesar PDF: {str(e)}")
            raise
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extraer texto de un archivo Word."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # También extraer de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extrayendo DOCX: {e}")
            self.errors.append(f"Error al procesar Word: {str(e)}")
            raise
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extraer texto de un archivo de texto plano."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Intentar con otra codificación
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error leyendo TXT: {e}")
            self.errors.append(f"Error al leer archivo de texto: {str(e)}")
            raise
    
    def parse_job_profile(self, text: str) -> Dict[str, Any]:
        """
        Parsea el texto de un perfil de cargo en datos estructurados.
        
        Algoritmo:
        1. Limpiar texto (quitar headers/footers)
        2. Identificar título (primera línea no vacía o encabezado)
        3. Buscar secciones por keywords
        4. Extraer bullet points bajo cada sección
        5. Parsear datos estructurados (key: value)
        6. Normalizar formato
        """
        # Limpiar texto
        text = self._clean_text(text)
        
        # Identificar título del cargo
        role_title = self._extract_role_title(text)
        
        # Identificar secciones
        sections = self._identify_sections(text)
        
        # Parsear cada sección
        result = {
            "role_title": role_title,
            "objective": self._parse_objective(sections.get("objetivo", "")),
            "hierarchy": self._parse_hierarchy(sections.get("estructura", "")),
            "responsibilities": self._parse_responsibilities(
                sections.get("responsabilidades", ""),
                text  # También buscar en todo el texto
            ),
            "requirements": self._parse_requirements(sections.get("perfil", "")),
            "skills": self._parse_skills(
                sections.get("conocimientos", ""),
                sections.get("competencias", ""),
                sections.get("perfil", "")
            ),
            "disc_profile": self._parse_disc(sections.get("disc", "")),
            "tools": self._parse_tools(sections.get("herramientas", "")),
            "conditions": self._parse_conditions(sections.get("condiciones", "")),
            "metadata": {
                "extraction_errors": self.errors,
                "extraction_warnings": self.warnings,
                "sections_found": list(sections.keys()),
            }
        }
        
        # Agregar descripción combinada
        result["description"] = self._build_description(result)
        
        return result
    
    def _clean_text(self, text: str) -> str:
        """Limpia el texto de headers, footers y caracteres extraños."""
        # Normalizar saltos de línea
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        
        # Quitar múltiples espacios
        text = re.sub(r' +', ' ', text)
        
        # Quitar líneas de numeración de página común
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        text = re.sub(r'\n\s*P[áa]gina\s*\d+\s*\n', '\n', text, flags=re.IGNORECASE)
        
        # Quitar URLs comunes de footers
        text = re.sub(r'https?://\S+', '', text)
        
        # Quitar emails comunes de footers
        text = re.sub(r'\S+@\S+\.\S+', '', text)
        
        # Normalizar líneas vacías múltiples
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def _extract_role_title(self, text: str) -> str:
        """Extrae el título del cargo (primera línea significativa o encabezado)."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        if not lines:
            self.warnings.append("No se encontró título del cargo")
            return ""
        
        # Buscar patrones de título comunes
        title_patterns = [
            r'^([A-Z][A-Za-z\s]+(?:Manager|Director|Engineer|Analyst|Coordinator|Specialist|Lead|Head|Chief|Officer|Consultant|Assistant|Representative|Supervisor|Developer|Designer|Architect))',
            r'^Perfil\s+de\s+Cargo[\s:]+([^\n]+)',
            r'^Job\s+Profile[\s:]+([^\n]+)',
        ]
        
        for pattern in title_patterns:
            match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
            if match:
                title = match.group(1).strip()
                # Limpiar metadatos del título
                title = re.sub(r'\|.*$', '', title).strip()
                title = re.sub(r'\s+Version\s+\d+.*$', '', title, flags=re.IGNORECASE).strip()
                return title
        
        # Si no hay patrón, tomar la primera línea significativa
        first_line = lines[0]
        
        # Limpiar metadatos comunes
        first_line = re.sub(r'\|.*$', '', first_line).strip()
        first_line = re.sub(r'NGDS\s*\|', '', first_line).strip()
        first_line = re.sub(r'Version\s+\d+.*$', '', first_line, flags=re.IGNORECASE).strip()
        
        # Si la primera línea es muy corta, combinar con la segunda
        if len(first_line) < 10 and len(lines) > 1:
            first_line = f"{first_line} {lines[1]}"
        
        return first_line
    
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identifica las secciones del documento por sus keywords."""
        sections = {}
        lines = text.split('\n')
        
        # Encontrar líneas que marcan el inicio de secciones
        section_boundaries = []
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Quitar caracteres especiales para matching
            line_clean = re.sub(r'[^\w\s]', '', line_lower)
            
            for section_key, keywords in self.SECTIONS.items():
                for keyword in keywords:
                    # Buscar keyword al inicio de línea o como título
                    if line_clean.startswith(keyword) or re.search(rf'^{keyword}[:\s]', line_clean):
                        section_boundaries.append((i, section_key, line))
                        break
        
        # Ordenar por posición
        section_boundaries.sort(key=lambda x: x[0])
        
        # Extraer contenido de cada sección
        for i, (start_idx, section_key, header) in enumerate(section_boundaries):
            # Encontrar dónde termina esta sección
            if i + 1 < len(section_boundaries):
                end_idx = section_boundaries[i + 1][0]
            else:
                end_idx = len(lines)
            
            # Extraer contenido (sin la línea del header)
            content_lines = lines[start_idx + 1:end_idx]
            content = '\n'.join(line for line in content_lines if line.strip())
            
            # Si ya existe esta sección, concatenar
            if section_key in sections:
                sections[section_key] += '\n' + content
            else:
                sections[section_key] = content
        
        return sections
    
    def _parse_objective(self, text: str) -> str:
        """Parsea la sección de objetivo del rol."""
        if not text:
            # Buscar en el texto general
            return ""
        
        # Limpiar y devolver el párrafo
        text = text.strip()
        
        # Quitar bullet points si es un solo párrafo
        if '\n' not in text and text.startswith(('-', '•', '*')):
            text = text[1:].strip()
        
        return text
    
    def _parse_hierarchy(self, text: str) -> Dict[str, str]:
        """Parsea la sección de estructura/jerarquía."""
        hierarchy = {
            "reports_to": "",
            "direct_reports": "",
            "level": "",
            "work_mode": "",
            "location": ""
        }
        
        if not text:
            return hierarchy
        
        for field, patterns in self.HIERARCHY_PATTERNS.items():
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    value = match.group(1).strip()
                    # Limpiar valor
                    value = re.sub(r'\s+', ' ', value)
                    hierarchy[field] = value
                    break
        
        return hierarchy
    
    def _parse_responsibilities(self, text: str, full_text: str = "") -> Dict[str, List[str]]:
        """
        Parsea la sección de responsabilidades.
        Soporta agrupación por categorías (ej: Financieras, Tesorería, etc.)
        """
        responsibilities = {}
        
        if not text:
            text = full_text
        
        if not text:
            return responsibilities
        
        # Buscar categorías seguidas de dos puntos
        category_pattern = r'^([A-Z][A-Za-z\s]+):\s*$'
        lines = text.split('\n')
        
        current_category = "General"
        current_items = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Verificar si es una categoría
            category_match = re.match(category_pattern, line)
            if category_match and not line.startswith(('-', '•', '*', '1.', '2.')):
                # Guardar categoría anterior
                if current_items:
                    responsibilities[current_category] = current_items
                
                current_category = category_match.group(1).strip()
                current_items = []
            else:
                # Es un item de responsabilidad
                item = self._clean_bullet_point(line)
                if item and len(item) > 5:  # Evitar líneas muy cortas
                    current_items.append(item)
        
        # Guardar última categoría
        if current_items:
            responsibilities[current_category] = current_items
        
        # Si no se encontraron categorías, buscar bullets generales
        if not responsibilities:
            bullets = self._extract_bullets(text)
            if bullets:
                responsibilities["Responsabilidades"] = bullets
        
        return responsibilities
    
    def _parse_requirements(self, text: str) -> Dict[str, str]:
        """Parsea la sección de requisitos/perfil requerido."""
        requirements = {
            "education": "",
            "experience_years": "",
            "experience_details": ""
        }
        
        if not text:
            return requirements
        
        # Buscar campos específicos
        for field, patterns in self.REQUIREMENTS_PATTERNS.items():
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
                if match:
                    value = match.group(1).strip()
                    value = re.sub(r'\s+', ' ', value)
                    requirements[field] = value
                    break
        
        # Extraer detalles adicionales de experiencia
        experience_details = []
        lines = text.split('\n')
        
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in [
                'experiencia en', 'experiencia previa', 'conocimiento en',
                'indispensable', 'deseable', 'requisito'
            ]):
                item = self._clean_bullet_point(line)
                if item:
                    experience_details.append(item)
        
        if experience_details:
            requirements["experience_details"] = "\n".join(experience_details)
        
        return requirements
    
    def _parse_skills(self, technical_text: str, soft_text: str, profile_text: str) -> Dict[str, List[str]]:
        """Parsea las secciones de conocimientos técnicos y competencias."""
        skills = {
            "technical": [],
            "soft": []
        }
        
        # Procesar conocimientos técnicos
        if technical_text:
            skills["technical"] = self._extract_bullets(technical_text)
        
        # Procesar competencias blandas
        if soft_text:
            skills["soft"] = self._extract_bullets(soft_text)
        
        # Si no se encontraron skills, buscar en el perfil general
        if not skills["technical"] and not skills["soft"] and profile_text:
            # Buscar sección de conocimientos en el perfil
            knowledge_match = re.search(
                r'(?:Conocimientos?|Skills?)[\s\w]*:\s*\n?((?:\s*[-•*\d][^\n]*\n?)+)',
                profile_text,
                re.IGNORECASE
            )
            if knowledge_match:
                skills["technical"] = self._extract_bullets(knowledge_match.group(1))
        
        return skills
    
    def _parse_disc(self, text: str) -> str:
        """Parsea la sección de perfil DISC."""
        if not text:
            return ""
        
        text = text.strip()
        
        # Buscar patrones DISC comunes
        disc_patterns = [
            r'(Concienzudo\s*\+\s*\w+\s*\([^)]+\))',
            r'(Dominante\s*\+\s*\w+\s*\([^)]+\))',
            r'(Influyente\s*\+\s*\w+\s*\([^)]+\))',
            r'(Estable\s*\+\s*\w+\s*\([^)]+\))',
            r'Perfil\s+DISC:\s*([^\n]+)',
        ]
        
        for pattern in disc_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        # Si no hay patrón específico, devolver el texto limpio
        return text
    
    def _parse_tools(self, text: str) -> List[str]:
        """Parsea la sección de herramientas/software."""
        if not text:
            return []
        
        return self._extract_bullets(text)
    
    def _parse_conditions(self, text: str) -> Dict[str, str]:
        """Parsea la sección de condiciones laborales."""
        conditions = {
            "salary_range": "",
            "benefits": "",
            "location": "",
            "work_schedule": ""
        }
        
        if not text:
            return conditions
        
        # Buscar patrones comunes
        patterns = {
            "salary_range": [
                r'salario[\s:]+([^\n]+)',
                r'remuneraci[oó]n[\s:]+([^\n]+)',
                r'salary[\s:]+([^\n]+)',
            ],
            "benefits": [
                r'beneficios[\s:]+([^\n]+(?:\n[^\n]+)*?)(?=\n\w|$)',
            ],
            "location": [
                r'ubicaci[oó]n[\s:]+([^\n]+)',
                r'lugar[\s:]+([^\n]+)',
            ],
            "work_schedule": [
                r'horario[\s:]+([^\n]+)',
                r'jornada[\s:]+([^\n]+)',
            ]
        }
        
        for field, field_patterns in patterns.items():
            for pattern in field_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    conditions[field] = match.group(1).strip()
                    break
        
        return conditions
    
    def _extract_bullets(self, text: str) -> List[str]:
        """Extrae items de una lista con bullets (•, -, *, números)."""
        items = []
        
        lines = text.split('\n')
        
        for line in lines:
            item = self._clean_bullet_point(line)
            if item and len(item) > 3:  # Evitar items muy cortos
                items.append(item)
        
        return items
    
    def _clean_bullet_point(self, line: str) -> str:
        """Limpia un bullet point de sus marcadores."""
        line = line.strip()
        
        # Quitar bullets comunes
        line = re.sub(r'^[\s•\-\*\+]+\s*', '', line)
        line = re.sub(r'^\d+[\.\)\-]\s*', '', line)
        
        # Quitar espacios extra
        line = re.sub(r'\s+', ' ', line)
        
        return line.strip()
    
    def _build_description(self, data: Dict[str, Any]) -> str:
        """Construye una descripción completa combinando objetivo y responsabilidades."""
        parts = []
        
        # Agregar objetivo
        if data.get("objective"):
            parts.append("## Objetivo del Rol\n")
            parts.append(data["objective"])
            parts.append("")
        
        # Agregar responsabilidades
        if data.get("responsabilities"):
            parts.append("\n## Responsabilidades\n")
            for category, items in data["responsabilities"].items():
                if len(data["responsabilities"]) > 1:
                    parts.append(f"\n**{category}:**")
                for item in items:
                    parts.append(f"- {item}")
        
        return "\n".join(parts)


# Función helper para uso directo
def extract_job_profile(file_path: str) -> Dict[str, Any]:
    """
    Función helper para extraer perfil de cargo desde un archivo.
    
    Args:
        file_path: Ruta al archivo (PDF, DOCX, DOC, TXT)
        
    Returns:
        Dict con los datos estructurados del perfil
    """
    extractor = JobProfileExtractor()
    return extractor.extract_from_document(file_path)