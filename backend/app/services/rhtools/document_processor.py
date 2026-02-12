"""Servicio de procesamiento de documentos - Extracción de texto."""
import io
import logging
import mimetypes
import os
import tempfile
import time
from pathlib import Path
from typing import Optional, Dict, Any, Tuple

import httpx

from app.core.config import settings
from app.models.rhtools import Document, DocumentTextExtraction, DocumentStatus, DocumentType

logger = logging.getLogger(__name__)


class DocumentProcessorError(Exception):
    """Error en procesamiento de documentos."""
    pass


class UnsupportedFileError(DocumentProcessorError):
    """Tipo de archivo no soportado."""
    pass


class ExtractionError(DocumentProcessorError):
    """Error al extraer texto."""
    pass


class DocumentProcessor:
    """Procesador de documentos para extracción de texto."""
    
    # MIME types soportados
    SUPPORTED_MIME_TYPES = {
        # PDFs
        'application/pdf': 'pdf',
        'application/x-pdf': 'pdf',
        # Word
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc',
        # Imágenes
        'image/jpeg': 'image',
        'image/jpg': 'image',
        'image/png': 'image',
        'image/tiff': 'image',
        'image/bmp': 'image',
        'image/webp': 'image',
        # Texto plano
        'text/plain': 'txt',
        'text/markdown': 'txt',
        'text/rtf': 'rtf',
    }
    
    def __init__(self, db_session=None, s3_client=None):
        """Inicializa el procesador.
        
        Args:
            db_session: Sesión de base de datos SQLAlchemy
            s3_client: Cliente S3 (boto3 o compatible)
        """
        self.db = db_session
        self.s3 = s3_client
        self._pdf_extractor = None
        self._docx_extractor = None
        self._ocr_engine = None
    
    def _get_pdf_extractor(self):
        """Lazy load del extractor de PDFs."""
        if self._pdf_extractor is None:
            try:
                # Intentar usar pdfplumber primero (mejor extracción)
                import pdfplumber
                self._pdf_extractor = 'pdfplumber'
                logger.info("Using pdfplumber for PDF extraction")
            except ImportError:
                # Fallback a PyPDF2
                try:
                    import PyPDF2
                    self._pdf_extractor = 'pypdf2'
                    logger.info("Using PyPDF2 for PDF extraction")
                except ImportError:
                    raise DocumentProcessorError(
                        "No PDF library available. Install pdfplumber or PyPDF2"
                    )
        return self._pdf_extractor
    
    def _get_docx_extractor(self):
        """Lazy load del extractor de Word."""
        if self._docx_extractor is None:
            try:
                import docx
                self._docx_extractor = 'python-docx'
                logger.info("Using python-docx for Word extraction")
            except ImportError:
                raise DocumentProcessorError(
                    "python-docx not available. Install: pip install python-docx"
                )
        return self._docx_extractor
    
    def _get_ocr_engine(self):
        """Lazy load del motor OCR."""
        if self._ocr_engine is None:
            try:
                import pytesseract
                from PIL import Image
                self._ocr_engine = 'tesseract'
                logger.info("Using Tesseract OCR")
            except ImportError:
                logger.warning("Tesseract not available, OCR will use AWS Textract fallback")
                self._ocr_engine = 'textract'
        return self._ocr_engine
    
    def detect_mime_type(self, file_path: str, filename: str) -> str:
        """Detecta el MIME type de un archivo.
        
        Args:
            file_path: Ruta al archivo
            filename: Nombre original del archivo
            
        Returns:
            MIME type detectado
        """
        # Intentar por extensión primero
        mime_type, _ = mimetypes.guess_type(filename)
        
        if mime_type:
            return mime_type
        
        # Fallback: leer magic bytes
        try:
            import magic
            mime_type = magic.from_file(file_path, mime=True)
            return mime_type
        except ImportError:
            pass
        
        # Fallback final por extensión
        ext = Path(filename).suffix.lower()
        mime_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.tiff': 'image/tiff',
            '.txt': 'text/plain',
            '.md': 'text/markdown',
        }
        return mime_map.get(ext, 'application/octet-stream')
    
    def is_supported(self, mime_type: str) -> bool:
        """Verifica si el MIME type es soportado.
        
        Args:
            mime_type: MIME type a verificar
            
        Returns:
            True si es soportado
        """
        return mime_type.lower() in self.SUPPORTED_MIME_TYPES
    
    def get_file_type(self, mime_type: str) -> str:
        """Obtiene el tipo de archivo simplificado.
        
        Args:
            mime_type: MIME type
            
        Returns:
            Tipo simplificado: pdf, docx, image, txt
        """
        return self.SUPPORTED_MIME_TYPES.get(mime_type.lower(), 'unknown')
    
    async def extract_text(self, document_id: str) -> DocumentTextExtraction:
        """Extrae texto de un documento.
        
        Args:
            document_id: ID del documento en la BD
            
        Returns:
            DocumentTextExtraction con el texto extraído
        """
        from sqlalchemy.ext.asyncio import AsyncSession
        
        if not self.db:
            raise DocumentProcessorError("Database session required")
        
        # Obtener documento
        from sqlalchemy import select
        result = await self.db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        
        if not document:
            raise DocumentProcessorError(f"Document {document_id} not found")
        
        # Actualizar estado
        document.status = DocumentStatus.PROCESSING.value
        await self.db.commit()
        
        start_time = time.time()
        
        try:
            # Descargar archivo de S3 o filesystem
            file_path = await self._get_document_file(document)
            
            # Detectar tipo
            mime_type = document.mime_type or self.detect_mime_type(file_path, document.original_filename)
            file_type = self.get_file_type(mime_type)
            
            if not self.is_supported(mime_type):
                raise UnsupportedFileError(f"Unsupported file type: {mime_type}")
            
            # Extraer texto según tipo
            extraction_result = await self._extract_by_type(file_path, file_type)
            
            duration_ms = int((time.time() - start_time) * 1000)
            
            # Crear registro de extracción
            extraction = DocumentTextExtraction(
                document_id=document.id,
                extracted_text=extraction_result['text'],
                text_length=len(extraction_result['text']),
                extraction_method=extraction_result['method'],
                extraction_duration_ms=duration_ms,
                ocr_confidence=extraction_result.get('ocr_confidence'),
                ocr_engine=extraction_result.get('ocr_engine')
            )
            
            self.db.add(extraction)
            
            # Actualizar documento
            document.status = DocumentStatus.PROCESSED.value
            document.processed_at = datetime.utcnow()
            await self.db.commit()
            await self.db.refresh(extraction)
            
            logger.info(
                f"Extracted {len(extraction_result['text'])} chars from "
                f"document {document_id} in {duration_ms}ms"
            )
            
            return extraction
            
        except Exception as e:
            document.status = DocumentStatus.ERROR.value
            await self.db.commit()
            
            logger.error(f"Error extracting text from {document_id}: {e}")
            raise ExtractionError(f"Extraction failed: {str(e)}")
        
        finally:
            # Limpiar archivo temporal si existe
            if 'file_path' in locals() and file_path.startswith(tempfile.gettempdir()):
                try:
                    os.remove(file_path)
                except:
                    pass
    
    async def _get_document_file(self, document: Document) -> str:
        """Obtiene el archivo del documento.
        
        Args:
            document: Documento
            
        Returns:
            Ruta al archivo local
        """
        if document.s3_bucket and document.s3_key and self.s3:
            # Descargar de S3
            temp_file = tempfile.NamedTemporaryFile(delete=False)
            try:
                self.s3.download_fileobj(
                    document.s3_bucket,
                    document.s3_key,
                    temp_file
                )
                temp_file.close()
                return temp_file.name
            except Exception as e:
                temp_file.close()
                os.unlink(temp_file.name)
                raise DocumentProcessorError(f"Failed to download from S3: {e}")
        
        # Usar filesystem local
        local_path = os.path.join(
            settings.UPLOAD_DIR,
            document.stored_filename
        )
        
        if not os.path.exists(local_path):
            raise DocumentProcessorError(f"File not found: {local_path}")
        
        return local_path
    
    async def _extract_by_type(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Extrae texto según el tipo de archivo.
        
        Args:
            file_path: Ruta al archivo
            file_type: Tipo de archivo (pdf, docx, image, txt)
            
        Returns:
            Dict con 'text', 'method', y opcionalmente 'ocr_confidence', 'ocr_engine'
        """
        if file_type == 'pdf':
            return await self._extract_pdf(file_path)
        elif file_type == 'docx':
            return await self._extract_docx(file_path)
        elif file_type == 'image':
            return await self._extract_image(file_path)
        elif file_type == 'txt':
            return await self._extract_txt(file_path)
        else:
            raise UnsupportedFileError(f"Extraction not implemented for type: {file_type}")
    
    async def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extrae texto de un PDF."""
        extractor = self._get_pdf_extractor()
        
        if extractor == 'pdfplumber':
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return {
                'text': '\n\n'.join(text_parts),
                'method': 'pdfplumber'
            }
        else:
            import PyPDF2
            
            text_parts = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text())
            
            return {
                'text': '\n\n'.join(text_parts),
                'method': 'pypdf2'
            }
    
    async def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extrae texto de un documento Word."""
        self._get_docx_extractor()
        import docx
        
        doc = docx.Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # También extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return {
            'text': '\n\n'.join(text_parts),
            'method': 'python-docx'
        }
    
    async def _extract_image(self, file_path: str) -> Dict[str, Any]:
        """Extrae texto de una imagen usando OCR."""
        ocr_engine = self._get_ocr_engine()
        
        if ocr_engine == 'tesseract':
            import pytesseract
            from PIL import Image
            
            image = Image.open(file_path)
            
            # Configurar para mejorar OCR de documentos
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(image, config=custom_config)
            
            # Obtener confianza
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            return {
                'text': text,
                'method': 'tesseract',
                'ocr_confidence': avg_confidence / 100,  # Normalizar a 0-1
                'ocr_engine': 'tesseract'
            }
        else:
            # Fallback a AWS Textract
            return await self._extract_with_textract(file_path)
    
    async def _extract_with_textract(self, file_path: str) -> Dict[str, Any]:
        """Extrae texto usando AWS Textract."""
        try:
            import boto3
            
            client = boto3.client('textract')
            
            with open(file_path, 'rb') as f:
                response = client.detect_document_text(
                    Document={'Bytes': f.read()}
                )
            
            text_parts = []
            for item in response['Blocks']:
                if item['BlockType'] == 'LINE':
                    text_parts.append(item['Text'])
            
            # Calcular confianza promedio
            confidences = [
                item['Confidence'] 
                for item in response['Blocks'] 
                if 'Confidence' in item
            ]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            return {
                'text': '\n'.join(text_parts),
                'method': 'aws_textract',
                'ocr_confidence': avg_confidence / 100,
                'ocr_engine': 'aws_textract'
            }
            
        except Exception as e:
            logger.error(f"Textract failed: {e}")
            raise ExtractionError(f"OCR failed: {e}")
    
    async def _extract_txt(self, file_path: str) -> Dict[str, Any]:
        """Extrae texto de un archivo de texto plano."""
        # Detectar encoding
        import chardet
        
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            detected = chardet.detect(raw_data)
            encoding = detected['encoding'] or 'utf-8'
        
        with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
            text = f.read()
        
        return {
            'text': text,
            'method': 'text_reader'
        }
    
    async def extract_text_from_bytes(
        self, 
        file_bytes: bytes, 
        filename: str,
        mime_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """Extrae texto directamente de bytes.
        
        Útil para procesamiento sin guardar en BD primero.
        
        Args:
            file_bytes: Contenido del archivo
            filename: Nombre del archivo
            mime_type: MIME type (opcional, se detecta si no se proporciona)
            
        Returns:
            Dict con 'text', 'method', etc.
        """
        # Guardar en archivo temporal
        suffix = Path(filename).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            if not mime_type:
                mime_type = self.detect_mime_type(tmp_path, filename)
            
            file_type = self.get_file_type(mime_type)
            return await self._extract_by_type(tmp_path, file_type)
            
        finally:
            os.unlink(tmp_path)
